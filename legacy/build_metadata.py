"""
Build per-application-area survey summary reports with historical weather,
Garmin unit metadata, processing specifications, and datum details.

Outputs:
  - survey_reports/<area_name>.docx : one Word report per area
  - survey_weather_summary.csv      : flat table, one row per outing
"""

from __future__ import annotations

import ast
import csv
import json
import math
import re
import time
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from functools import lru_cache
from pathlib import Path
from urllib.parse import urlencode
from urllib.request import urlopen
from zoneinfo import ZoneInfo

from docx import Document
from pyproj import CRS
from shapely.geometry import shape


REPO_DIR = Path(__file__).resolve().parent
APPLICATION_AREAS_PATH = REPO_DIR / ""
MATCHED_TRACKS_PATH = Path(r"")
REPORTS_DIR = REPO_DIR / "survey_reports"
SUMMARY_CSV_PATH = REPO_DIR / "survey_weather_summary.csv"
MOSAIC_CONFIG_PATH = REPO_DIR / "garmin_mosaic.py"
AREA_MOSAICS_DIR = MATCHED_TRACKS_PATH.parent / "application_area_mosaics"

LOCAL_TZ = ZoneInfo("America/Chicago")
PER_FILE_DURATION_MIN = 30  # approximate length of each RSD survey run
GPS_DATUM_NAME = "WGS 84 (EPSG:4326)"
FEATURE_DETECTION_NOTE = (
    "Automated feature detection is not performed by this workflow; findings in "
    "the report are limited to collection conditions, available depth/range "
    "metadata, and the presence of processed intensity mosaics for manual review."
)

ARCHIVE_URL = "https://archive-api.open-meteo.com/v1/archive"
MARINE_URL = "https://marine-api.open-meteo.com/v1/marine"

WEATHER_HOURLY = [
    "temperature_2m",
    "wind_speed_10m",
    "wind_gusts_10m",
    "wind_direction_10m",
    "precipitation",
    "cloud_cover",
]
MARINE_HOURLY = [
    "wave_height",
    "wave_period",
    "wave_direction",
    "wind_wave_height",
    "swell_wave_height",
]
PROCESSING_FIELDS = {
    "OUTPUT_RESOLUTION",
    "MAX_RANGE_FALLBACK",
    "PORT_MAX_RANGE_OVERRIDE_M",
    "PAYLOAD_MODE_OVERRIDE_PORT",
    "PAYLOAD_MODE_OVERRIDE_STARBOARD",
    "PORT_CHANNEL_OVERRIDE_ID",
    "HEADING_SMOOTH_WINDOW",
    "MIN_SPEED_MS",
}

FILENAME_RE = re.compile(r"^(\d{2})([A-Z]{3})(\d{2})-(\d{2})(\d{2})-\d+", re.IGNORECASE)
MONTHS = {
    "JAN": 1, "FEB": 2, "MAR": 3, "APR": 4, "MAY": 5, "JUN": 6,
    "JUL": 7, "AUG": 8, "SEP": 9, "OCT": 10, "NOV": 11, "DEC": 12,
}


def parse_filename_datetime(file_name: str) -> datetime | None:
    match = FILENAME_RE.match(file_name)
    if not match:
        return None
    day, mon, yr, hh, mm = match.groups()
    month = MONTHS.get(mon.upper())
    if month is None:
        return None
    year = 2000 + int(yr)
    return datetime(year, month, int(day), int(hh), int(mm), tzinfo=LOCAL_TZ)


def sanitize_name(value: str):
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return safe.strip("._") or "unnamed"


def safe_float(value):
    if value in (None, "", "nan", "NaN"):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def safe_int(value):
    if value in (None, "", "nan", "NaN"):
        return None
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return None


def dedupe_preserve_order(values):
    seen = set()
    ordered = []
    for value in values:
        if value in (None, "", []):
            continue
        if value in seen:
            continue
        seen.add(value)
        ordered.append(value)
    return ordered


def fmt(value, suffix=""):
    if value is None:
        return "n/a"
    return f"{value}{suffix}"


def avg_or_none(values):
    clean = [v for v in values if v is not None]
    if not clean:
        return None
    return round(sum(clean) / len(clean), 2)


def min_or_none(values):
    clean = [v for v in values if v is not None]
    return round(min(clean), 2) if clean else None


def max_or_none(values):
    clean = [v for v in values if v is not None]
    return round(max(clean), 2) if clean else None


def range_text(min_value, max_value, suffix=""):
    if min_value is None and max_value is None:
        return "n/a"
    if min_value == max_value:
        return fmt(min_value, suffix)
    return f"{fmt(min_value, suffix)} to {fmt(max_value, suffix)}"


def load_application_areas():
    payload = json.loads(APPLICATION_AREAS_PATH.read_text(encoding="utf-8"))
    areas = {}
    for feature in payload.get("features", []):
        props = feature.get("properties", {}) or {}
        name = props.get("Name")
        geom = feature.get("geometry")
        if not name or not geom:
            continue
        poly = shape(geom)
        if poly.is_empty:
            continue
        centroid = poly.centroid
        areas[name] = {
            "name": name,
            "safe_name": sanitize_name(name),
            "collector": props.get("FolderPath") or "",
            "status": props.get("PopupInfo") or "",
            "object_id": props.get("OBJECTID"),
            "centroid_lon": float(centroid.x),
            "centroid_lat": float(centroid.y),
            "area_m2": float(props.get("Shape_Area") or 0.0),
        }
    return areas


def get_output_paths(rsd_file: Path):
    output_base_dir = rsd_file.parent / f"garmin_output_{rsd_file.stem}"
    meta_dir = output_base_dir / "meta"
    processed_dir = output_base_dir / "processed"
    return output_base_dir, meta_dir, processed_dir


@lru_cache(maxsize=None)
def load_processing_config():
    source = MOSAIC_CONFIG_PATH.read_text(encoding="utf-8")
    tree = ast.parse(source, filename=str(MOSAIC_CONFIG_PATH))
    config = {}
    for node in tree.body:
        if not isinstance(node, ast.Assign) or len(node.targets) != 1:
            continue
        target = node.targets[0]
        if not isinstance(target, ast.Name) or target.id not in PROCESSING_FIELDS:
            continue
        try:
            config[target.id] = ast.literal_eval(node.value)
        except Exception:
            continue
    return config


def read_first_csv_row(path: Path):
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8", newline="") as fh:
        return next(csv.DictReader(fh), {}) or {}


def summarize_all_meta(path: Path):
    summary = {
        "depth_mean_m": None,
        "depth_min_m": None,
        "depth_max_m": None,
        "range_mean_m": None,
        "range_min_m": None,
        "range_max_m": None,
        "utm_zones": [],
    }
    if not path.exists():
        return summary

    depth_values = []
    range_values = []
    utm_zones = set()

    with path.open("r", encoding="utf-8", newline="") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            depth = safe_float(row.get("inst_dep_m"))
            if depth is not None:
                depth_values.append(depth)

            ping_range = safe_float(row.get("max_range"))
            if ping_range is not None:
                range_values.append(ping_range)

            zone = safe_int(row.get("utm_zone"))
            if zone is not None:
                utm_zones.add(zone)

    summary["depth_mean_m"] = avg_or_none(depth_values)
    summary["depth_min_m"] = min_or_none(depth_values)
    summary["depth_max_m"] = max_or_none(depth_values)
    summary["range_mean_m"] = avg_or_none(range_values)
    summary["range_min_m"] = min_or_none(range_values)
    summary["range_max_m"] = max_or_none(range_values)
    summary["utm_zones"] = sorted(utm_zones)
    return summary


@lru_cache(maxsize=None)
def load_track_support_metadata(rsd_path_str: str):
    rsd_file = Path(rsd_path_str)
    _, meta_dir, processed_dir = get_output_paths(rsd_file)
    dat_meta = read_first_csv_row(meta_dir / "DAT_meta.csv")
    all_meta_summary = summarize_all_meta(meta_dir / "All-Garmin-Sonar-MetaData.csv")
    intensity_path = processed_dir / "intensity.tif"

    return {
        "dat_meta_path": meta_dir / "DAT_meta.csv",
        "all_meta_path": meta_dir / "All-Garmin-Sonar-MetaData.csv",
        "intensity_path": intensity_path,
        "has_intensity_path": intensity_path.exists(),
        "unit_product_number": dat_meta.get("unit_product_number") or "",
        "unit_id_type": dat_meta.get("unit_id_type") or "",
        "unit_software_version": dat_meta.get("unit_software_version") or "",
        "channel_count": safe_int(dat_meta.get("channel_count")),
        "depth_mean_m": all_meta_summary["depth_mean_m"],
        "depth_min_m": all_meta_summary["depth_min_m"],
        "depth_max_m": all_meta_summary["depth_max_m"],
        "range_mean_m": all_meta_summary["range_mean_m"],
        "range_min_m": all_meta_summary["range_min_m"],
        "range_max_m": all_meta_summary["range_max_m"],
        "utm_zones": all_meta_summary["utm_zones"],
    }


def load_matched_tracks():
    payload = json.loads(MATCHED_TRACKS_PATH.read_text(encoding="utf-8"))
    tracks = []
    for feature in payload.get("features", []):
        props = feature.get("properties", {}) or {}
        file_name = props.get("file_name") or ""
        file_path_value = props.get("file_path")
        if not file_name or not file_path_value:
            continue
        dt = parse_filename_datetime(file_name)
        if dt is None:
            continue
        rsd_file = Path(file_path_value).expanduser().resolve()
        support = load_track_support_metadata(str(rsd_file))
        tracks.append(
            {
                "file_name": file_name,
                "file_path": rsd_file,
                "start_local": dt,
                "end_local": dt + timedelta(minutes=PER_FILE_DURATION_MIN),
                "area_names": list(props.get("application_area_names") or []),
                "support": support,
            }
        )
    return tracks


def group_outings(tracks, areas):
    outings = defaultdict(list)
    for track in tracks:
        local_date = track["start_local"].date()
        for area_name in track["area_names"]:
            if area_name not in areas:
                continue
            outings[(area_name, local_date)].append(track)
    return outings


def fetch_json(url: str, params: dict, retries: int = 3, backoff: float = 2.0):
    qs = urlencode(params, doseq=True)
    full_url = f"{url}?{qs}"
    last_exc = None
    for attempt in range(retries):
        try:
            with urlopen(full_url, timeout=45) as resp:
                return json.loads(resp.read().decode("utf-8"))
        except Exception as exc:
            last_exc = exc
            if attempt + 1 < retries:
                time.sleep(backoff * (attempt + 1))
    raise RuntimeError(f"Request failed: {full_url} ({last_exc})")


_weather_cache: dict = {}
_marine_cache: dict = {}


def cache_key(lat: float, lon: float, start_date: str, end_date: str):
    return (round(lat, 2), round(lon, 2), start_date, end_date)


def fetch_weather(lat: float, lon: float, start_date: str, end_date: str):
    key = cache_key(lat, lon, start_date, end_date)
    if key not in _weather_cache:
        _weather_cache[key] = fetch_json(
            ARCHIVE_URL,
            {
                "latitude": round(lat, 4),
                "longitude": round(lon, 4),
                "start_date": start_date,
                "end_date": end_date,
                "hourly": ",".join(WEATHER_HOURLY),
                "wind_speed_unit": "kn",
                "temperature_unit": "fahrenheit",
                "precipitation_unit": "inch",
                "timezone": "UTC",
            },
        )
    return _weather_cache[key]


def fetch_marine(lat: float, lon: float, start_date: str, end_date: str):
    key = cache_key(lat, lon, start_date, end_date)
    if key not in _marine_cache:
        _marine_cache[key] = fetch_json(
            MARINE_URL,
            {
                "latitude": round(lat, 4),
                "longitude": round(lon, 4),
                "start_date": start_date,
                "end_date": end_date,
                "hourly": ",".join(MARINE_HOURLY),
                "length_unit": "imperial",
                "timezone": "UTC",
            },
        )
    return _marine_cache[key]


def _parse_iso_utc(ts: str) -> datetime:
    return datetime.fromisoformat(ts).replace(tzinfo=timezone.utc)


def average_over_window(payload: dict, variables: list[str], start_utc: datetime, end_utc: datetime):
    hourly = (payload or {}).get("hourly") or {}
    times = hourly.get("time") or []
    if not times:
        return {v: None for v in variables}

    results = {}
    indices = [i for i, t in enumerate(times) if start_utc <= _parse_iso_utc(t) <= end_utc]
    if not indices:
        mid = start_utc + (end_utc - start_utc) / 2
        indices = [min(range(len(times)), key=lambda i: abs(_parse_iso_utc(times[i]) - mid))]

    for var in variables:
        series = hourly.get(var) or []
        vals = [series[i] for i in indices if i < len(series) and series[i] is not None]
        if not vals:
            results[var] = None
            continue
        if var in {"wind_direction_10m", "wave_direction"}:
            sins = sum(math.sin(math.radians(v)) for v in vals)
            coss = sum(math.cos(math.radians(v)) for v in vals)
            deg = math.degrees(math.atan2(sins, coss))
            if deg < 0:
                deg += 360
            results[var] = round(deg, 0)
        else:
            results[var] = round(sum(vals) / len(vals), 2)
    return results


def compass_from_deg(deg):
    if deg is None:
        return ""
    dirs = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
            "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
    idx = int((deg / 22.5) + 0.5) % 16
    return dirs[idx]


def utm_epsg_from_zone(zone: int, lat: float):
    return 32600 + zone if lat >= 0 else 32700 + zone


def projected_crs_text(area, tracks):
    zones = sorted({zone for track in tracks for zone in track["support"].get("utm_zones", [])})
    if not zones:
        zone = int((area["centroid_lon"] + 180.0) / 6.0) + 1
        zones = [min(max(zone, 1), 60)]

    crs_texts = []
    for zone in zones:
        epsg = utm_epsg_from_zone(zone, area["centroid_lat"])
        crs_texts.append(f"WGS 84 / UTM zone {zone}{'N' if area['centroid_lat'] >= 0 else 'S'} (EPSG:{epsg})")
    return ", ".join(crs_texts)


def summarize_track_set(area, tracks):
    product_numbers = dedupe_preserve_order(
        [track["support"].get("unit_product_number") for track in tracks]
    )
    unit_ids = dedupe_preserve_order(
        [track["support"].get("unit_id_type") for track in tracks]
    )
    software_versions = dedupe_preserve_order(
        [track["support"].get("unit_software_version") for track in tracks]
    )
    channel_counts = dedupe_preserve_order(
        [str(track["support"].get("channel_count")) for track in tracks if track["support"].get("channel_count") is not None]
    )

    depth_means = [track["support"].get("depth_mean_m") for track in tracks]
    depth_mins = [track["support"].get("depth_min_m") for track in tracks]
    depth_maxs = [track["support"].get("depth_max_m") for track in tracks]
    range_means = [track["support"].get("range_mean_m") for track in tracks]
    range_mins = [track["support"].get("range_min_m") for track in tracks]
    range_maxs = [track["support"].get("range_max_m") for track in tracks]

    area_mosaic_path = AREA_MOSAICS_DIR / area["safe_name"] / f"{area['safe_name']}_intensity_clipped.tif"

    return {
        "unit_make": "Garmin",
        "unit_product_numbers": product_numbers,
        "unit_id_types": unit_ids,
        "unit_software_versions": software_versions,
        "channel_counts": channel_counts,
        "depth_mean_m": avg_or_none(depth_means),
        "depth_min_m": min_or_none(depth_mins),
        "depth_max_m": max_or_none(depth_maxs),
        "range_mean_m": avg_or_none(range_means),
        "range_min_m": min_or_none(range_mins),
        "range_max_m": max_or_none(range_maxs),
        "gps_datum": GPS_DATUM_NAME,
        "projected_crs": projected_crs_text(area, tracks),
        "track_count_with_mosaic": sum(1 for track in tracks if track["support"].get("has_intensity_path")),
        "area_mosaic_path": area_mosaic_path,
        "area_mosaic_exists": area_mosaic_path.exists(),
    }


def build_outing_row(area, local_date, tracks, processing_config):
    start_local = min(t["start_local"] for t in tracks)
    end_local = max(t["end_local"] for t in tracks)
    start_utc = start_local.astimezone(timezone.utc)
    end_utc = end_local.astimezone(timezone.utc)
    start_date = start_utc.date().isoformat()
    end_date = end_utc.date().isoformat()

    weather_payload = fetch_weather(area["centroid_lat"], area["centroid_lon"], start_date, end_date)
    marine_payload = fetch_marine(area["centroid_lat"], area["centroid_lon"], start_date, end_date)
    weather = average_over_window(weather_payload, WEATHER_HOURLY, start_utc, end_utc)
    marine = average_over_window(marine_payload, MARINE_HOURLY, start_utc, end_utc)
    track_summary = summarize_track_set(area, tracks)

    return {
        "area_name": area["name"],
        "collector": area["collector"],
        "status": area["status"],
        "object_id": area["object_id"],
        "centroid_lat": round(area["centroid_lat"], 5),
        "centroid_lon": round(area["centroid_lon"], 5),
        "local_date": local_date.isoformat(),
        "start_local": start_local.strftime("%Y-%m-%d %H:%M %Z"),
        "end_local": end_local.strftime("%Y-%m-%d %H:%M %Z"),
        "file_count": len(tracks),
        "file_names": ",".join(sorted(t["file_name"] for t in tracks)),
        "temperature_f": weather.get("temperature_2m"),
        "wind_speed_kt": weather.get("wind_speed_10m"),
        "wind_gust_kt": weather.get("wind_gusts_10m"),
        "wind_dir_deg": weather.get("wind_direction_10m"),
        "wind_dir_compass": compass_from_deg(weather.get("wind_direction_10m")),
        "precip_in": weather.get("precipitation"),
        "cloud_cover_pct": weather.get("cloud_cover"),
        "wave_height_ft": marine.get("wave_height"),
        "wave_period_s": marine.get("wave_period"),
        "wave_dir_deg": marine.get("wave_direction"),
        "wave_dir_compass": compass_from_deg(marine.get("wave_direction")),
        "wind_wave_height_ft": marine.get("wind_wave_height"),
        "swell_wave_height_ft": marine.get("swell_wave_height"),
        "unit_make": track_summary["unit_make"],
        "unit_product_numbers": ",".join(track_summary["unit_product_numbers"]) or "n/a",
        "unit_id_types": ",".join(track_summary["unit_id_types"]) or "n/a",
        "unit_software_versions": ",".join(track_summary["unit_software_versions"]) or "n/a",
        "channel_counts": ",".join(track_summary["channel_counts"]) or "n/a",
        "gps_datum": track_summary["gps_datum"],
        "projected_crs": track_summary["projected_crs"],
        "depth_mean_m": track_summary["depth_mean_m"],
        "depth_range_m": range_text(track_summary["depth_min_m"], track_summary["depth_max_m"], " m"),
        "range_mean_m": track_summary["range_mean_m"],
        "range_range_m": range_text(track_summary["range_min_m"], track_summary["range_max_m"], " m"),
        "output_resolution_m": processing_config.get("OUTPUT_RESOLUTION"),
        "fallback_range_m": processing_config.get("MAX_RANGE_FALLBACK"),
        "port_override_range_m": processing_config.get("PORT_MAX_RANGE_OVERRIDE_M"),
        "heading_smooth_window_pings": processing_config.get("HEADING_SMOOTH_WINDOW"),
        "min_speed_ms": processing_config.get("MIN_SPEED_MS"),
        "payload_mode_port": processing_config.get("PAYLOAD_MODE_OVERRIDE_PORT"),
        "payload_mode_starboard": processing_config.get("PAYLOAD_MODE_OVERRIDE_STARBOARD"),
        "port_channel_override_id": processing_config.get("PORT_CHANNEL_OVERRIDE_ID"),
        "area_intensity_mosaic_path": str(track_summary["area_mosaic_path"]) if track_summary["area_mosaic_exists"] else "",
    }


def write_summary_csv(rows, path: Path):
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    fieldnames = list(rows[0].keys())
    with path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def add_labeled_paragraph(doc: Document, label: str, value):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(str(value) if value not in (None, "") else "n/a")


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def value_range_across_rows(rows, key):
    values = [safe_float(row.get(key)) for row in rows]
    values = [value for value in values if value is not None]
    if not values:
        return "n/a"
    return range_text(round(min(values), 2), round(max(values), 2))


def build_findings_summary(area, rows_for_area, tracks_for_area, track_summary):
    dates = sorted({row["local_date"] for row in rows_for_area})
    sentences = [
        (
            f"{area['name']} was surveyed across {len(rows_for_area)} outing(s) and "
            f"{len(tracks_for_area)} Garmin RSD file(s) on {', '.join(dates)}."
        )
    ]

    if track_summary["depth_mean_m"] is not None:
        sentences.append(
            "Available ping metadata indicate an average instrument depth of "
            f"{fmt(track_summary['depth_mean_m'], ' m')} with a range of "
            f"{range_text(track_summary['depth_min_m'], track_summary['depth_max_m'], ' m')}."
        )

    if track_summary["range_mean_m"] is not None:
        sentences.append(
            "Available sonar metadata indicate an average per-ping range of "
            f"{fmt(track_summary['range_mean_m'], ' m')} with a range of "
            f"{range_text(track_summary['range_min_m'], track_summary['range_max_m'], ' m')}."
        )

    if rows_for_area:
        sentences.append(
            "Observed collection conditions across the surveyed outings ranged from "
            f"{value_range_across_rows(rows_for_area, 'wind_speed_kt')} kt wind, "
            f"{value_range_across_rows(rows_for_area, 'wave_height_ft')} ft wave height, and "
            f"{value_range_across_rows(rows_for_area, 'precip_in')} in rainfall."
        )

    if track_summary["area_mosaic_exists"]:
        sentences.append(
            f"A clipped processed intensity mosaic is available at {track_summary['area_mosaic_path']}."
        )
    else:
        sentences.append(
            "A clipped processed intensity mosaic for this area has not yet been generated."
        )

    sentences.append(FEATURE_DETECTION_NOTE)
    return " ".join(sentences)


def write_area_report_docx(area, rows_for_area, tracks_for_area, processing_config, out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    track_summary = summarize_track_set(area, tracks_for_area)
    report_path = out_dir / f"{area['safe_name']}.docx"

    doc = Document()
    doc.add_heading(f"{area['name']} Survey Summary and Metadata", level=0)

    add_labeled_paragraph(doc, "Who collected the data", area["collector"] or "n/a")
    add_labeled_paragraph(doc, "Status", area["status"] or "n/a")
    add_labeled_paragraph(doc, "OBJECTID", area["object_id"])
    add_labeled_paragraph(doc, "Area centroid", f"{area['centroid_lat']:.5f}, {area['centroid_lon']:.5f}")
    add_labeled_paragraph(doc, "Horizontal datum for GPS coordinates", track_summary["gps_datum"])
    add_labeled_paragraph(doc, "Projected CRS for processed mosaics", track_summary["projected_crs"])

    doc.add_heading("Equipment and Processing", level=1)
    add_bullet(doc, f"Equipment make: {track_summary['unit_make']}")
    add_bullet(doc, f"Garmin product number(s): {', '.join(track_summary['unit_product_numbers']) or 'n/a'}")
    add_bullet(doc, f"Garmin unit ID type(s): {', '.join(track_summary['unit_id_types']) or 'n/a'}")
    add_bullet(doc, f"Garmin software version(s): {', '.join(track_summary['unit_software_versions']) or 'n/a'}")
    add_bullet(doc, f"Recorded channel count(s): {', '.join(track_summary['channel_counts']) or 'n/a'}")
    add_bullet(doc, f"Output mosaic resolution: {fmt(processing_config.get('OUTPUT_RESOLUTION'), ' m/pixel')}")
    add_bullet(doc, f"Average available ping range: {fmt(track_summary['range_mean_m'], ' m')}")
    add_bullet(doc, f"Range span from metadata: {range_text(track_summary['range_min_m'], track_summary['range_max_m'], ' m')}")
    add_bullet(doc, f"Fallback range setting: {fmt(processing_config.get('MAX_RANGE_FALLBACK'), ' m')}")
    add_bullet(doc, f"Port override range setting: {fmt(processing_config.get('PORT_MAX_RANGE_OVERRIDE_M'), ' m')}")
    add_bullet(doc, f"Heading smoothing window: {fmt(processing_config.get('HEADING_SMOOTH_WINDOW'), ' pings')}")
    add_bullet(doc, f"Minimum speed filter: {fmt(processing_config.get('MIN_SPEED_MS'), ' m/s')}")
    add_bullet(doc, f"Payload mode (port/starboard): {fmt(processing_config.get('PAYLOAD_MODE_OVERRIDE_PORT'))} / {fmt(processing_config.get('PAYLOAD_MODE_OVERRIDE_STARBOARD'))}")
    add_bullet(doc, f"Port channel override: {fmt(processing_config.get('PORT_CHANNEL_OVERRIDE_ID'))}")
    add_bullet(doc, "Operating frequency: not exposed as a human-readable field in the current Garmin RSD metadata export")
    add_bullet(doc, "Transect spacing: not automatically derived in the current workflow")

    doc.add_heading("Processed Mosaic Deliverable", level=1)
    if track_summary["area_mosaic_exists"]:
        add_labeled_paragraph(doc, "Processed and mosaicked SONAR image (GeoTIFF)", track_summary["area_mosaic_path"])
    else:
        add_labeled_paragraph(doc, "Processed and mosaicked SONAR image (GeoTIFF)", "Not yet generated")

    if not rows_for_area:
        doc.add_heading("Summary of Findings", level=1)
        doc.add_paragraph("No Garmin RSD surveys have been recorded for this area.")
        doc.save(report_path)
        return

    doc.add_heading("Survey Timing and Conditions", level=1)
    add_labeled_paragraph(doc, "Survey dates", ", ".join(sorted({row['local_date'] for row in rows_for_area})))
    add_labeled_paragraph(doc, "Total outings", len(rows_for_area))
    add_labeled_paragraph(doc, "Total RSD files", len(tracks_for_area))

    table = doc.add_table(rows=1, cols=8)
    headers = table.rows[0].cells
    header_labels = [
        "Date", "Local window", "Files", "Wind (kt)", "Wave ht (ft)",
        "Rain (in)", "Cloud (%)", "Temp (F)",
    ]
    for cell, label in zip(headers, header_labels):
        cell.text = label

    for row in sorted(rows_for_area, key=lambda item: item["local_date"]):
        cells = table.add_row().cells
        cells[0].text = row["local_date"]
        cells[1].text = f"{row['start_local'].split(' ')[1]}-{row['end_local'].split(' ')[1]}"
        cells[2].text = str(row["file_count"])
        cells[3].text = fmt(row["wind_speed_kt"])
        cells[4].text = fmt(row["wave_height_ft"])
        cells[5].text = fmt(row["precip_in"])
        cells[6].text = fmt(row["cloud_cover_pct"])
        cells[7].text = fmt(row["temperature_f"])

    doc.add_heading("Summary of Findings", level=1)
    doc.add_paragraph(build_findings_summary(area, rows_for_area, tracks_for_area, track_summary))

    doc.add_heading("RSD Files by Outing", level=1)
    for row in sorted(rows_for_area, key=lambda item: item["local_date"]):
        doc.add_paragraph(row["local_date"], style="List Bullet")
        for file_name in row["file_names"].split(","):
            doc.add_paragraph(file_name, style="List Bullet 2")

    doc.add_paragraph(
        "Weather source: Open-Meteo ERA5 archive. Marine source: Open-Meteo wave model. "
        "Values are mean values over the interpreted outing window; direction values use a circular mean."
    )
    doc.save(report_path)


def main():
    if not APPLICATION_AREAS_PATH.exists():
        raise SystemExit(f"Missing: {APPLICATION_AREAS_PATH}")
    if not MATCHED_TRACKS_PATH.exists():
        raise SystemExit(f"Missing: {MATCHED_TRACKS_PATH}")
    if not MOSAIC_CONFIG_PATH.exists():
        raise SystemExit(f"Missing: {MOSAIC_CONFIG_PATH}")

    processing_config = load_processing_config()
    areas = load_application_areas()
    tracks = load_matched_tracks()
    print(f"Loaded {len(areas)} application areas, {len(tracks)} matched RSD tracks")

    outings = group_outings(tracks, areas)
    print(f"Grouped into {len(outings)} outings (unique area+date combos)")

    all_rows = []
    rows_by_area = defaultdict(list)
    tracks_by_area = defaultdict(dict)

    total = len(outings)
    for i, ((area_name, local_date), outing_tracks) in enumerate(sorted(outings.items()), start=1):
        area = areas[area_name]
        print(f"[{i}/{total}] {area_name} {local_date} ({len(outing_tracks)} file(s))")
        try:
            row = build_outing_row(area, local_date, outing_tracks, processing_config)
        except Exception as exc:
            print(f"  ! fetch failed: {exc}")
            continue
        all_rows.append(row)
        rows_by_area[area_name].append(row)
        for track in outing_tracks:
            tracks_by_area[area_name][str(track["file_path"]).lower()] = track

    write_summary_csv(all_rows, SUMMARY_CSV_PATH)
    print(f"Wrote {SUMMARY_CSV_PATH}")

    for name, area in sorted(areas.items()):
        area_tracks = list(tracks_by_area.get(name, {}).values())
        write_area_report_docx(area, rows_by_area.get(name, []), area_tracks, processing_config, REPORTS_DIR)
    print(f"Wrote {len(areas)} area report(s) to {REPORTS_DIR}")


if __name__ == "__main__":
    main()
